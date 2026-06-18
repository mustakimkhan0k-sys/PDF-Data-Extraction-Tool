import os
import pdfplumber
import pandas as pd
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

# ------------------------------
# PDF Extraction Function
# ------------------------------
def extract_pdf_data(pdf_path, start_page, end_page, output_format, output_path):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            if end_page == 0 or end_page > total_pages:
                end_page = total_pages

            extracted_data = []
            for page_num in range(start_page - 1, end_page):
                text = pdf.pages[page_num].extract_text() or ""
                extracted_data.append({
                    "Page": page_num + 1,
                    "Content": text.strip()
                })

        if output_format == "Excel (.xlsx)":
            df = pd.DataFrame(extracted_data)
            df.to_excel(output_path, index=False)
        elif output_format == "Word (.docx)":
            doc = Document()
            for entry in extracted_data:
                doc.add_heading(f"Page {entry['Page']}", level=2)
                doc.add_paragraph(entry["Content"])
                doc.add_page_break()
            doc.save(output_path)
        elif output_format == "Text (.txt)":
            with open(output_path, "w", encoding="utf-8") as f:
                for entry in extracted_data:
                    f.write(f"--- Page {entry['Page']} ---\n{entry['Content']}\n\n")
        else:
            messagebox.showerror("Error", "Unsupported output format selected.")
            return

        messagebox.showinfo("Success", f"Data extracted successfully!\nSaved to:\n{output_path}")

    except Exception as e:
        messagebox.showerror("Error", f"An error occurred:\n{str(e)}")


# ------------------------------
# GUI Functions
# ------------------------------
def browse_pdf():
    path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
    pdf_path_var.set(path)

def choose_output_file():
    filetypes = []
    if output_format_var.get() == "Excel (.xlsx)":
        filetypes = [("Excel Files", "*.xlsx")]
    elif output_format_var.get() == "Word (.docx)":
        filetypes = [("Word Files", "*.docx")]
    elif output_format_var.get() == "Text (.txt)":
        filetypes = [("Text Files", "*.txt")]

    path = filedialog.asksaveasfilename(defaultextension=filetypes[0][1], filetypes=filetypes)
    output_path_var.set(path)

def run_extraction():
    pdf_path = pdf_path_var.get()
    start_page = int(start_page_var.get() or 1)
    end_page = int(end_page_var.get() or 0)
    output_format = output_format_var.get()
    output_path = output_path_var.get()

    if not pdf_path:
        messagebox.showerror("Error", "Please select a PDF file.")
        return
    if not output_path:
        messagebox.showerror("Error", "Please select an output file.")
        return

    extract_pdf_data(pdf_path, start_page, end_page, output_format, output_path)


# ------------------------------
# Main GUI Setup
# ------------------------------
root = tk.Tk()
root.title("PDF → Extractor (select PDF, pages, output format)")

# 🖥️ Make Fullscreen
root.state('zoomed')  # Works on Windows (maximized)
# root.attributes('-fullscreen', True)  # Alternative for true fullscreen (press ESC to exit)

# Variables
pdf_path_var = tk.StringVar()
start_page_var = tk.StringVar()
end_page_var = tk.StringVar()
output_format_var = tk.StringVar(value="Excel (.xlsx)")
output_path_var = tk.StringVar()

# Layout Frame
frame = ttk.Frame(root, padding=30)
frame.pack(fill="both", expand=True)

# Widgets
ttk.Label(frame, text="Source PDF:", font=("Segoe UI", 12)).grid(row=0, column=0, sticky="w", pady=10)
ttk.Entry(frame, textvariable=pdf_path_var, width=80).grid(row=0, column=1, pady=10, padx=10)
ttk.Button(frame, text="Browse...", command=browse_pdf).grid(row=0, column=2, pady=10, padx=5)

ttk.Label(frame, text="Start page (1-based):", font=("Segoe UI", 12)).grid(row=1, column=0, sticky="w", pady=10)
ttk.Entry(frame, textvariable=start_page_var, width=10).grid(row=1, column=1, sticky="w", pady=10)
ttk.Label(frame, text="End page (inclusive):", font=("Segoe UI", 12)).grid(row=1, column=1, sticky="e", pady=10)
ttk.Entry(frame, textvariable=end_page_var, width=10).grid(row=1, column=2, sticky="w", pady=10)

ttk.Label(frame, text="Output format:", font=("Segoe UI", 12)).grid(row=2, column=0, sticky="w", pady=10)
ttk.Combobox(frame, textvariable=output_format_var,
             values=["Excel (.xlsx)", "Word (.docx)", "Text (.txt)"],
             width=20, state="readonly").grid(row=2, column=1, sticky="w", pady=10)

ttk.Label(frame, text="Output file:", font=("Segoe UI", 12)).grid(row=3, column=0, sticky="w", pady=10)
ttk.Entry(frame, textvariable=output_path_var, width=80).grid(row=3, column=1, pady=10, padx=10)
ttk.Button(frame, text="Choose...", command=choose_output_file).grid(row=3, column=2, pady=10, padx=5)

# Buttons
ttk.Button(frame, text="Run Extraction", command=run_extraction).grid(row=4, column=1, pady=30, sticky="e", padx=10)
ttk.Button(frame, text="Quit", command=root.quit).grid(row=4, column=2, pady=30, sticky="w", padx=10)

# Status label
status_label = ttk.Label(frame, text="Ready", font=("Segoe UI", 10))
status_label.grid(row=5, column=0, columnspan=3, sticky="w", pady=10)

root.mainloop()
